from docxtpl import DocxTemplate
from docx import Document
from typing import Dict, Optional
from pathlib import Path
import docx2txt
import shutil

from .file_type_base import FileType
from .context_helpers import escape_chars
from ..constants import REMOVE_SENTINEL


def _remove_sentinel_paragraphs(docx_path: Path) -> int:
    """
    Scan *docx_path* for any paragraph whose text contains REMOVE_SENTINEL,
    remove those paragraphs from the document XML, and save in-place.

    Covers both body paragraphs and paragraphs inside table cells.
    Returns the number of paragraphs removed (0 = nothing to do).
    """
    doc = Document(str(docx_path))
    removed = 0

    # Body paragraphs
    for para in list(doc.paragraphs):
        if REMOVE_SENTINEL in para.text:
            para._element.getparent().remove(para._element)
            removed += 1

    # Paragraphs inside table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in list(cell.paragraphs):
                    if REMOVE_SENTINEL in para.text:
                        para._element.getparent().remove(para._element)
                        removed += 1

    if removed:
        doc.save(str(docx_path))
        print(f"Sentinel removal: {removed} paragraph(s) containing '{REMOVE_SENTINEL}' removed.")

    return removed


class DOCXf(FileType):
    def get_resume_str(self) -> str:
        return docx2txt.process(self.res_path)

    def post_llm_process(self, context: Dict[str, str], metadata: Optional[dict] = None) -> Path:
        self.context = escape_chars(context, "docx")

        orig = self.res_path
        working_copy = self._build_output_path(metadata)
        shutil.copy2(orig, working_copy)

        doc = DocxTemplate(working_copy)
        doc.init_docx()

        doc.render(self.context)
        try:
            doc.save(working_copy)
        except OSError as e:
            print(f"Render failed: {e}")
        else:
            print(f"Render good.")

        _remove_sentinel_paragraphs(working_copy)

        return working_copy
